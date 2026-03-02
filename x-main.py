"""
X (Twitter) Research AI Agent — FastAPI Backend
Stack: FastAPI + RapidAPI Twitter135 + Gemini 1.5 Pro + python-docx
"""

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional
import httpx
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json, os, uuid, base64
from datetime import datetime
from pathlib import Path

app = FastAPI(title="X Research Agent")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

RAPIDAPI_KEY  = os.getenv("RAPIDAPI_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

TWITTER_HOST = "twitter135.p.rapidapi.com"
TWITTER_BASE = "https://twitter135.p.rapidapi.com"

projects_db: dict = {}
analyses_db: dict = {}
briefs_db:   dict = {}

Path("outputs").mkdir(exist_ok=True)


# ─────────────────────────────────────────
# MODELS
# ─────────────────────────────────────────

class Project(BaseModel):
    name: str
    client: str
    brand_bible: str
    brief_template: str

class SearchRequest(BaseModel):
    project_id: str
    keyword: str
    count: int = 20
    content_type: str = "mixed"   # "tweets" | "threads" | "mixed"

class AnalyzeRequest(BaseModel):
    project_id: str
    tweet_ids: list[str]

class BriefRequest(BaseModel):
    project_id: str
    analysis_ids: list[str]


# ─────────────────────────────────────────
# PROJECTS
# ─────────────────────────────────────────

@app.post("/projects")
async def create_project(p: Project):
    pid = str(uuid.uuid4())[:8]
    projects_db[pid] = {**p.dict(), "id": pid, "created_at": datetime.now().isoformat()}
    return projects_db[pid]

@app.get("/projects")
async def list_projects():
    return list(projects_db.values())

@app.get("/projects/{pid}")
async def get_project(pid: str):
    if pid not in projects_db:
        raise HTTPException(404, "Project not found")
    return projects_db[pid]

@app.put("/projects/{pid}")
async def update_project(pid: str, p: Project):
    if pid not in projects_db:
        raise HTTPException(404, "Project not found")
    projects_db[pid] = {**p.dict(), "id": pid, "created_at": projects_db[pid]["created_at"]}
    return projects_db[pid]


# ─────────────────────────────────────────
# RAPIDAPI HEADERS
# ─────────────────────────────────────────

def rapidapi_headers():
    return {
        "X-RapidAPI-Key": RAPIDAPI_KEY,
        "X-RapidAPI-Host": TWITTER_HOST,
    }


# ─────────────────────────────────────────
# X SEARCH  (Twitter135 on RapidAPI)
# ─────────────────────────────────────────

def parse_tweet(t: dict) -> dict:
    """Normalize a raw tweet object from Twitter135 API."""
    legacy = t.get("legacy", t)
    user   = t.get("core", {}).get("user_results", {}).get("result", {}).get("legacy", {})

    # Detect thread: has_thread marker or is self-reply
    is_thread = bool(legacy.get("self_thread"))
    tweet_type = "thread" if is_thread else "tweet"

    # Detect quote tweet
    if t.get("quoted_status_result"):
        tweet_type = "quote_tweet"

    # Images
    media = legacy.get("entities", {}).get("media", [])
    image_urls = [m.get("media_url_https") for m in media if m.get("type") == "photo"]

    return {
        "id":           legacy.get("id_str") or t.get("rest_id", ""),
        "text":         legacy.get("full_text", ""),
        "author":       user.get("name", ""),
        "handle":       user.get("screen_name", ""),
        "author_followers": user.get("followers_count", 0),
        "verified":     user.get("verified", False) or user.get("is_blue_verified", False),
        "tweet_type":   tweet_type,
        "likes":        legacy.get("favorite_count", 0),
        "retweets":     legacy.get("retweet_count", 0),
        "replies":      legacy.get("reply_count", 0),
        "quotes":       legacy.get("quote_count", 0),
        "bookmarks":    legacy.get("bookmark_count", 0),
        "image_urls":   image_urls,
        "has_images":   len(image_urls) > 0,
        "created_at":   legacy.get("created_at", ""),
        "tweet_url":    f"https://x.com/{user.get('screen_name','')}/status/{legacy.get('id_str','')}",
    }


@app.post("/search")
async def search_x(req: SearchRequest):
    """Search X via Twitter135 RapidAPI."""
    headers = rapidapi_headers()

    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(
            f"{TWITTER_BASE}/v2/SearchTimeline/",
            headers=headers,
            params={
                "query": req.keyword,
                "count": req.count,
                "type":  "Latest",
            }
        )

    if resp.status_code != 200:
        raise HTTPException(502, f"Twitter API error {resp.status_code}: {resp.text[:300]}")

    raw = resp.json()

    # Navigate the Twitter GraphQL response tree
    instructions = (
        raw.get("data", {})
           .get("search_by_raw_query", {})
           .get("search_timeline", {})
           .get("timeline", {})
           .get("instructions", [])
    )

    tweets = []
    for instruction in instructions:
        for entry in instruction.get("entries", []):
            content = entry.get("content", {})
            item_content = content.get("itemContent", {})
            tweet_result = item_content.get("tweet_results", {}).get("result", {})
            if tweet_result:
                try:
                    tweets.append(parse_tweet(tweet_result))
                except Exception:
                    pass

    # Filter by content type
    if req.content_type == "threads":
        tweets = [t for t in tweets if t["tweet_type"] == "thread"]
    elif req.content_type == "tweets":
        tweets = [t for t in tweets if t["tweet_type"] == "tweet"]

    # Sort by engagement (likes + retweets)
    tweets.sort(key=lambda t: t["likes"] + t["retweets"] * 3, reverse=True)

    return {"keyword": req.keyword, "tweets": tweets[:req.count], "count": len(tweets)}


# ─────────────────────────────────────────
# FETCH REPLIES
# ─────────────────────────────────────────

async def fetch_replies(tweet_id: str, handle: str, limit: int = 20) -> list[str]:
    """Fetch top replies to a tweet."""
    try:
        async with httpx.AsyncClient(timeout=20) as client:
            resp = await client.get(
                f"{TWITTER_BASE}/v2/TweetDetail/",
                headers=rapidapi_headers(),
                params={"id": tweet_id, "count": limit}
            )
        data = resp.json()
        instructions = (
            data.get("data", {})
                .get("threaded_conversation_with_injections_v2", {})
                .get("instructions", [])
        )
        replies = []
        for inst in instructions:
            for entry in inst.get("entries", []):
                for item in entry.get("content", {}).get("items", []):
                    tweet_res = item.get("item", {}).get("itemContent", {}).get("tweet_results", {}).get("result", {})
                    if tweet_res:
                        text = tweet_res.get("legacy", {}).get("full_text", "")
                        if text and not text.startswith("RT "):
                            replies.append(text)
        return replies[:limit]
    except Exception:
        return []


# ─────────────────────────────────────────
# GEMINI PROMPTS
# ─────────────────────────────────────────

HOOK_PROMPT = """
You are a senior creative strategist analyzing X (Twitter) posts for a brand marketing team.

Analyze this post and return structured JSON with EXACTLY these keys:

{
  "hook_type": "one of: Bold Claim | Contrarian Take | Story Opener | Stat Drop | Question | Curiosity Gap | List Bait | Hot Take | Personal Confession | Call Out",
  "hook_line": "the exact opening line (first sentence or clause)",
  "hook_score": integer 1-10 for how scroll-stopping the hook is,
  "hook_breakdown": "why this first line works — what psychological trigger it pulls",
  "format": "one of: Single Tweet | Thread | Quote Tweet | Tweet + Image | Tweet + Poll",
  "tone": "one of: Authoritative | Conversational | Provocative | Educational | Aspirational | Humorous | Vulnerable | Data-driven",
  "audience": "describe the exact person this is written for in one sentence",
  "writing_style": "describe the sentence structure, vocabulary level, rhythm — be specific",
  "emotional_trigger": "the core emotion being activated: curiosity / FOMO / validation / outrage / inspiration / humor / fear",
  "shareability": "High / Medium / Low",
  "shareability_reason": "one sentence on why people retweet or don't",
  "key_idea": "the core point or argument in one sentence",
  "content_angle": "the strategic angle: personal story / industry insight / counter-narrative / how-to / hot take / social proof",
  "visual_direction": "if there's an image or describe what visual would best accompany this tweet — be specific about style, mood, composition",
  "adaptation_hooks": ["3 specific hook rewrites a brand could use, inspired by this post's approach"]
}

Return ONLY valid JSON. No markdown, no explanation.
"""

REPLY_PROMPT = """
You are analyzing replies to a viral X (Twitter) post to extract audience intelligence for a creative strategist.

Replies:
{replies}

Return ONLY valid JSON with EXACTLY these keys:
{
  "top_reactions": ["the 3-5 most common types of reactions or sentiments expressed"],
  "audience_questions": ["real questions the audience is asking — direct quotes if striking"],
  "friction_points": ["objections, doubts, or pushback in the replies"],
  "amplifier_phrases": ["specific phrases or words the audience uses naturally that resonate"],
  "reply_sentiment": "Positive / Mixed / Critical / Polarized",
  "community_insight": "2-3 sentences on what this reply section reveals about the audience's worldview, language, and what they respond to"
}
"""

IMAGE_PROMPT = """
You are a creative director analyzing the visual accompanying this X (Twitter) post.

Analyze the image and return JSON with EXACTLY these keys:
{
  "visual_style": "describe the overall aesthetic in 5-8 words",
  "mood": "the emotional tone the image creates",
  "composition": "key compositional choices — what's foregrounded, color palette, lighting",
  "text_overlay": "any text in the image — transcribe it",
  "brand_cues": "colors, logos, or brand signals visible",
  "why_it_works": "one sentence on why this visual stops the scroll",
  "creative_direction": "3 specific visual directions a brand could take, inspired by this image"
}
Return ONLY valid JSON.
"""


# ─────────────────────────────────────────
# ANALYZE TWEETS
# ─────────────────────────────────────────

@app.post("/analyze")
async def analyze_tweets(req: AnalyzeRequest):
    results = []
    model = genai.GenerativeModel("gemini-1.5-pro")

    # We need full tweet data — the frontend passes tweet data as part of analyze
    # In production cache search results; here we re-fetch by tweet ID
    for tweet_id in req.tweet_ids:
        try:
            # Fetch tweet detail
            async with httpx.AsyncClient(timeout=20) as client:
                resp = await client.get(
                    f"{TWITTER_BASE}/v2/TweetDetail/",
                    headers=rapidapi_headers(),
                    params={"id": tweet_id, "count": 1}
                )
            raw = resp.json()

            # Extract tweet from detail response
            tweet_data = {}
            instructions = (
                raw.get("data", {})
                   .get("threaded_conversation_with_injections_v2", {})
                   .get("instructions", [])
            )
            for inst in instructions:
                for entry in inst.get("entries", []):
                    items = entry.get("content", {}).get("items", [])
                    item_content = entry.get("content", {}).get("itemContent", {})
                    tweet_res = item_content.get("tweet_results", {}).get("result", {})
                    if tweet_res and tweet_res.get("legacy", {}).get("id_str") == tweet_id:
                        tweet_data = parse_tweet(tweet_res)
                        break

            if not tweet_data:
                tweet_data = {"id": tweet_id, "text": "", "image_urls": [], "handle": ""}

            tweet_text = tweet_data.get("text", "")
            image_urls = tweet_data.get("image_urls", [])

            # ── Hook analysis ──
            hook_analysis = {}
            try:
                parts = [f"Tweet text:\n\n{tweet_text}\n\n{HOOK_PROMPT}"]

                # Add image if present
                if image_urls:
                    async with httpx.AsyncClient(timeout=15) as client:
                        img_resp = await client.get(image_urls[0])
                    if img_resp.status_code == 200:
                        img_b64 = base64.b64encode(img_resp.content).decode()
                        content_type = img_resp.headers.get("content-type", "image/jpeg")
                        parts = [
                            {"inline_data": {"mime_type": content_type, "data": img_b64}},
                            f"Tweet text:\n\n{tweet_text}\n\n{HOOK_PROMPT}"
                        ]

                hook_resp = model.generate_content(parts)
                hook_analysis = json.loads(hook_resp.text.strip().lstrip("```json").rstrip("```"))
            except Exception as e:
                hook_analysis = {"error": str(e)}

            # ── Image analysis ──
            image_analysis = {}
            if image_urls:
                try:
                    async with httpx.AsyncClient(timeout=15) as client:
                        img_resp = await client.get(image_urls[0])
                    img_b64 = base64.b64encode(img_resp.content).decode()
                    content_type = img_resp.headers.get("content-type", "image/jpeg")
                    img_resp2 = model.generate_content([
                        {"inline_data": {"mime_type": content_type, "data": img_b64}},
                        IMAGE_PROMPT
                    ])
                    image_analysis = json.loads(img_resp2.text.strip().lstrip("```json").rstrip("```"))
                except Exception as e:
                    image_analysis = {"error": str(e)}

            # ── Reply analysis ──
            reply_analysis = {}
            replies = await fetch_replies(tweet_id, tweet_data.get("handle", ""))
            if replies:
                try:
                    reply_text = "\n".join([f"- {r}" for r in replies])
                    reply_resp = model.generate_content(REPLY_PROMPT.format(replies=reply_text))
                    reply_analysis = json.loads(reply_resp.text.strip().lstrip("```json").rstrip("```"))
                except Exception as e:
                    reply_analysis = {"error": str(e)}

            analysis_id = str(uuid.uuid4())[:8]
            record = {
                "id":             analysis_id,
                "tweet_id":       tweet_id,
                "project_id":     req.project_id,
                "text":           tweet_text,
                "author":         tweet_data.get("author", ""),
                "handle":         tweet_data.get("handle", ""),
                "tweet_url":      tweet_data.get("tweet_url", ""),
                "tweet_type":     tweet_data.get("tweet_type", "tweet"),
                "likes":          tweet_data.get("likes", 0),
                "retweets":       tweet_data.get("retweets", 0),
                "replies_count":  tweet_data.get("replies", 0),
                "image_urls":     image_urls,
                "hook_analysis":  hook_analysis,
                "image_analysis": image_analysis,
                "reply_analysis": reply_analysis,
                "analyzed_at":    datetime.now().isoformat(),
            }
            analyses_db[analysis_id] = record
            results.append(record)

        except Exception as e:
            results.append({"tweet_id": tweet_id, "error": str(e)})

    return {"analyses": results}


@app.get("/analyses/{project_id}")
async def get_analyses(project_id: str):
    return [a for a in analyses_db.values() if a.get("project_id") == project_id]


# ─────────────────────────────────────────
# BRIEF GENERATION
# ─────────────────────────────────────────

BRIEF_PROMPT = """
You are a senior creative strategist at a top agency specializing in X (Twitter) content.

Generate a complete creative brief using the research below.

## Brand Bible
{brand_bible}

## Brief Template
{brief_template}

## Research: {post_count} analyzed X posts

{research_summary}

Instructions:
- Fill every section of the template with specific, actionable content
- Use exact hook lines and phrases from the research — not vague direction
- Reference the visual patterns you found across posts
- Write for someone who will brief a copywriter and a designer tomorrow
- Be ruthlessly specific. "Use a bold opening stat" is bad. "Open with a counter-intuitive data point like '83% of X posts get zero engagement — here's what the 17% do differently'" is good.

Return the completed brief using the template structure exactly.
"""


@app.post("/generate-brief")
async def generate_brief(req: BriefRequest):
    project = projects_db.get(req.project_id)
    if not project:
        raise HTTPException(404, "Project not found")

    selected = [analyses_db[aid] for aid in req.analysis_ids if aid in analyses_db]
    if not selected:
        raise HTTPException(400, "No valid analyses found")

    parts = []
    for a in selected:
        h = a.get("hook_analysis", {})
        r = a.get("reply_analysis", {})
        v = a.get("image_analysis", {})
        eng = f"♥ {a.get('likes',0):,}  ↺ {a.get('retweets',0):,}  💬 {a.get('replies_count',0):,}"

        part = f"""
POST (@{a.get('handle','?')} · {a.get('tweet_type','').upper()}) — {eng}
Text: {a.get('text','')[:280]}
Hook Type: {h.get('hook_type')} | Hook Line: "{h.get('hook_line')}" | Score: {h.get('hook_score')}/10
Hook Breakdown: {h.get('hook_breakdown')}
Tone: {h.get('tone')} | Audience: {h.get('audience')}
Writing Style: {h.get('writing_style')}
Emotional Trigger: {h.get('emotional_trigger')} | Shareability: {h.get('shareability')}
Shareability Reason: {h.get('shareability_reason')}
Content Angle: {h.get('content_angle')}
Visual Direction: {h.get('visual_direction')}
Adaptation Hooks: {' | '.join(h.get('adaptation_hooks', []))}
{f"Visual Style: {v.get('visual_style')} | Mood: {v.get('mood')}" if v else ""}
{f"Reply Sentiment: {r.get('reply_sentiment')} | Community: {r.get('community_insight')}" if r else ""}
{f"Amplifier Phrases: {', '.join(r.get('amplifier_phrases',[]))}" if r and r.get('amplifier_phrases') else ""}
"""
        parts.append(part)

    research_summary = "\n---\n".join(parts)
    model = genai.GenerativeModel("gemini-1.5-pro")
    resp = model.generate_content(
        BRIEF_PROMPT.format(
            brand_bible=project["brand_bible"],
            brief_template=project["brief_template"],
            post_count=len(selected),
            research_summary=research_summary,
        )
    )

    brief_id = str(uuid.uuid4())[:8]
    record = {
        "id":           brief_id,
        "project_id":   req.project_id,
        "project_name": project["name"],
        "client":       project["client"],
        "content":      resp.text,
        "analysis_ids": req.analysis_ids,
        "post_count":   len(selected),
        "created_at":   datetime.now().isoformat(),
    }
    briefs_db[brief_id] = record
    return record


@app.get("/briefs/{project_id}")
async def get_briefs(project_id: str):
    return [b for b in briefs_db.values() if b.get("project_id") == project_id]


# ─────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────

@app.get("/export/docx/{brief_id}")
async def export_docx(brief_id: str):
    brief = briefs_db.get(brief_id)
    if not brief:
        raise HTTPException(404, "Brief not found")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Title block
    h = doc.add_heading(f"{brief['client'].upper()} — X CONTENT BRIEF", 0)
    h.runs[0].font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    h.runs[0].font.size = Pt(22)

    meta = doc.add_paragraph(
        f"Project: {brief['project_name']}   ·   Based on {brief['post_count']} posts   ·   {brief['created_at'][:10]}"
    )
    meta.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    meta.runs[0].font.size = Pt(10)
    doc.add_paragraph("")

    for line in brief["content"].split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1D, 0x9B, 0xF0)  # X blue
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
        else:
            p = doc.add_paragraph()
            # Bold inline **text**
            import re
            parts = re.split(r"\*\*(.+?)\*\*", line)
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    doc.add_paragraph("")
    footer = doc.add_paragraph(f"X Research Agent · {datetime.now().strftime('%B %d, %Y')}")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    path = f"outputs/brief_{brief_id}.docx"
    doc.save(path)
    return FileResponse(
        path,
        filename=f"x_brief_{brief['client']}_{brief_id}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)
